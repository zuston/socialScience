#!coding:utf-8

'''
人民网官员信息抽取,实时监控信息更新
'''
import sys
import time
import os
sys.path.append("..")
from bs4 import BeautifulSoup as bs
from ztool import Http
import xlwt
from docx import Document
from docx.shared import Pt
from util import Mongo
import logging

mongoCollectionName = "renmin"
mongoDb = Mongo.Mongo().db[mongoCollectionName]

'''
bean类文件
'''
class bean(object):
    def __init__(self):
        self.name = ''
        self.position = ''
        self.sex = ''
        self.birth = ''
        self.nativePlace = ''
        self.education = ''
        self.resume = ''
        self.area = ''
        self.city = ''
        self.province = ''
        self.duty = ''

    def setPosition(self,position):
        self.position = position
    def setSex(self,sex):
        self.sex = sex
    def setBirth(self,birth):
        self.birth = birth
    def setNativePlace(self,nativePlace):
        self.nativePlace = nativePlace
    def setEducation(self,education):
        self.education = education
    def setResume(self,resume):
        self.resume = resume
    def setArea(self,area):
        self.area = area
    def setCity(self,city):
        self.city = city
    def setProvince(self,province):
        self.province = province
    def setName(self,name):
        self.name = name
    def setDuty(self,duty):
        self.duty = duty

'''
导入文件中工具
'''
class common(object):
    def __init__(self):
        pass

    @staticmethod
    def write2Word(fileName,list):
        length = len(list)
        dictInfo = {11:"县级",10:"副省级",9:"省级"}
        if(os.path.exists(fileName)):
            document = Document(fileName)
        else:
            document = Document()
            document.add_heading(unicode(dictInfo[length]+"领导信息"), 0)

        headerDict = {11:list[0]+" "+list[1]+" "+list[2],
                      10:list[0]+" "+list[1],
                      9:list[0]}
        document.add_heading(unicode(headerDict[length]),level=3)

        offset = 11-length

        p = document.add_paragraph("")
        p.add_run(unicode(list[3-offset])).italic = True
        p.add_run(unicode("    "+list[5-offset]))
        document.add_paragraph(unicode("生年:  "+list[7-offset]))
        document.add_paragraph(unicode("性别:  "+list[6-offset]))
        document.add_paragraph(unicode("籍贯:  "+list[8-offset]))
        document.add_paragraph(unicode("学历:  "+list[9-offset]))
        document.add_paragraph(unicode("简历:  "+list[10-offset]))
        document.add_paragraph(unicode(""))

        document.save(fileName)


'''
县级领导信息收集
'''
class county(object):
    def __init__(self):
        self.flag = 0
        self.baseUrl = "http://ldzl.people.com.cn/dfzlk/front/"
        self.excelFileName = './data/county.xls'
        self.startUrl = 'http://ldzl.people.com.cn/dfzlk/front/xian1.htm'
        self.wordFileName = './data/officerWord/county.docx'

    def start(self):
        http = Http.Http()
        code, msg, res = http.open(self.startUrl)
        print "爬取状态:" + str(code)
        print "爬取信息:" + msg
        self.parser(res)

    def parser(self,res):
        soup = bs(res, "html.parser")
        list = soup.find("ul", "brown clearfix").find_all("li")
        # 所有的dict值
        rdict = {}
        for one in list[:-3]:
            url = one.a['href']
            name = one.a['title'].split(" ")[1].strip()
            rdict[name] = url
            # print name+":"+url
        print "总计省份:" + str(len(list[:-3]))
        return self.getSecond(rdict)

    def getSecond(self,rdict):
        http = Http.Http()
        for key, value in rdict.items():
            print key + " " + value
            url = self.baseUrl + value
            code, msg, res = http.open(url)
            if (code == 200):
                self.parserData(key, res)

    def parserData(self,province, res):
        soup = bs(res, "html.parser")
        base = soup.find("div", "fr p2j_reports_right title_2j sjzlk")

        for cityName in base.find_all('h2', class_='this'):
            for city in cityName.find_next_sibling('div', 'zlk_list').find_all('ol', 'fl'):
                for oneLine in city.find_all('li', class_=False):
                    # 任职地区
                    areaName = oneLine.span.string
                    # 党委领导姓名
                    partyCommitteeName = oneLine.a.string
                    # 党委领导简历具体信息
                    partyCommitteeUrl = oneLine.a['href']

                    govList = oneLine.find_all('a')
                    # 政府领导姓名
                    govCommitteeName = govList[1].string
                    # 政府领导简历具体信息链接
                    govCommitteeUrl = govList[1]['href']

                    # print cityName.text + " " + areaName + " " + partyCommitteeName + " " + govCommitteeName

                    partyCommitteePostion = ''
                    partyCommitteeSex = ''
                    partyCommitteeBirth = ''
                    partyCommitteeNativePlace = ''
                    partyCommitteeEducation = ''
                    partyCommitteeResume = ''
                    partyPostionArea = areaName
                    partyPostionCity = cityName.text
                    partyProvince = province
                    if (partyCommitteeName != '空缺'):
                        partyRes = county.resumeInfo(partyCommitteeUrl,self.baseUrl)
                        if (len(partyRes) != 0):
                            partyCommitteePostion = partyRes[0]
                            partyCommitteeSex = partyRes[1]
                            partyCommitteeBirth = partyRes[2]
                            partyCommitteeNativePlace = partyRes[3]
                            partyCommitteeEducation = partyRes[4]
                            partyCommitteeResume = partyRes[5]

                    partyList = [partyProvince, partyPostionCity, partyPostionArea, partyCommitteeName, "党委领导",
                                 partyCommitteePostion, partyCommitteeSex, partyCommitteeBirth,
                                 partyCommitteeNativePlace, partyCommitteeEducation, partyCommitteeResume]


                    govCommitteePosition = ''
                    govCommitteeSex = ''
                    govCommitteeBirth = ''
                    govCommitteeNativePlace = ''
                    govCommitteeEducation = ''
                    govCommitteeResume = ''
                    govPostionArea = areaName
                    govPostionCity = cityName.text
                    govProvince = province
                    if (govCommitteeName != '空缺'):
                        govRes = county.resumeInfo(govCommitteeUrl,self.baseUrl)
                        if len(govRes) != 0:
                            govCommitteePosition = govRes[0]
                            govCommitteeSex = govRes[1]
                            govCommitteeBirth = govRes[2]
                            govCommitteeNativePlace = govRes[3]
                            govCommitteeEducation = govRes[4]
                            govCommitteeResume = govRes[5]

                    # print govPostionCity + '  区:' + govPostionArea + '  姓名:' + govCommitteeName + '  职位:' + govCommitteePosition + '  性别:' + govCommitteeSex + '  生年:' + govCommitteeBirth + '  籍贯:' + govCommitteeNativePlace + '  教育:' + govCommitteeEducation

                    govList = [govProvince, govPostionCity, govPostionArea, govCommitteeName, "政府领导",
                               govCommitteePosition, govCommitteeSex, govCommitteeBirth, govCommitteeNativePlace,
                               govCommitteeEducation, govCommitteeResume]
                    # self.flag = county.write2excel(self.flag,self.excelFileName,govList)
                    # self.flag = county.write2excel(self.flag,self.excelFileName,partyList)
                    # time.sleep(2)
                    # common.write2Word(self.wordFileName,govList)
                    # common.write2Word(self.wordFileName,partyList)

                    key = 3
                    partyList = [partyProvince, partyPostionCity, partyPostionArea, partyCommitteeName, "党委领导",
                                 partyCommitteePostion, partyCommitteeSex, partyCommitteeBirth,
                                 partyCommitteeNativePlace, partyCommitteeEducation, partyCommitteeResume]
                    officerDictParty = {
                        "province" : partyProvince,
                        "city" : partyPostionCity,
                        "area" : partyPostionArea,
                        "name" : partyCommitteeName,
                        "duty" : "党委书记",
                        "position" : partyCommitteePostion,
                        "sex" : partyCommitteeSex,
                        "birth" : partyCommitteeBirth,
                        "nativePlace" : partyCommitteeNativePlace,
                        "education" : partyCommitteeEducation,
                        "resume" : partyCommitteeResume
                    }
                    if self.filterMongo(officerDictParty,key):
                        logging.warn("[+]数据更新\t" + partyPostionArea + "\t" + partyCommitteePostion + "\t" + partyCommitteeName)
                        MongoDao.save(key,officerDictParty)
                    else:
                        logging.warn(
                            "[+]数据未更新\t" + partyPostionArea + "\t" + partyCommitteePostion + "\t" + partyCommitteeName)

                    officerDictGov = {
                        "province": govProvince,
                        "city": govPostionCity,
                        "area": govPostionArea,
                        "name": govCommitteeName,
                        "duty": "政府",
                        "position": govCommitteePosition,
                        "sex": govCommitteeSex,
                        "birth": govCommitteeBirth,
                        "nativePlace": govCommitteeNativePlace,
                        "education": govCommitteeEducation,
                        "resume": govCommitteeResume
                    }

                    if self.filterMongo(officerDictGov,key):
                        logging.warn("[+]数据更新\t" + govPostionArea + "\t" + govCommitteePosition + "\t" + govCommitteeName)

                        MongoDao.save(key,officerDictGov)
                    else:
                        logging.warn("[+]数据未更新\t" + govPostionArea + "\t" + govCommitteePosition + "\t" + govCommitteeName)




    def filterMongo(self, officerDict, key):
        return MongoDao.getData(officerDict, key)

    @staticmethod
    def resumeInfo(url,baseUrl):
        resList = list()
        url = baseUrl + url
        http = Http.Http()
        code, msg, res = http.open(url)
        if code != 200:
            return []
        soup = bs(res, "html.parser")
        postion = soup.find('span', class_='red2').string
        if postion==None:
            postion=''
        resList.append(postion)
        infoBase = soup.find('dd').find('p')
        infoBaseParser = str(infoBase).split('\n')[:-1]
        for line in infoBaseParser:
            value = line.split('</b>')[1].split('<br>')[0]
            if value==None:
                value=''
            resList.append(value)

        info = soup.find('div', class_='p2j_text')
        resList.append(str(info.text))
        return resList

    @staticmethod
    def write2excel(flag,excelFileName,list=None):
        if flag == 0:
            if (os.path.exists(excelFileName)):
                os.remove(excelFileName)
            excelWorkBook = xlwt.Workbook()
            sheet = excelWorkBook.add_sheet('officer')
            sheet.write(0, 0, u'省、直辖市')
            sheet.write(0, 1, u'市,直辖市城区')
            sheet.write(0, 2, u'县区')
            sheet.write(0, 3, u'姓名')
            sheet.write(0, 4, u'党委,政府')
            sheet.write(0, 5, u'职位')
            sheet.write(0, 6, u'性别')
            sheet.write(0, 7, u'生年')
            sheet.write(0, 8, u'籍贯')
            sheet.write(0, 9, u'学历')
            sheet.write(0, 10, u'简历')
            flag += 1
            excelWorkBook.save(excelFileName)
        else:
            import xlrd
            from xlutils.copy import copy
            rb = xlrd.open_workbook(excelFileName)
            wb = copy(rb)
            ws = wb.get_sheet(0)
            i = 0
            for data in list:
                ws.write(flag, i, unicode(data))
                i += 1
            flag += 1
            wb.save(excelFileName)
        return flag

    @staticmethod
    def write2Word(fileName,list):
        flag = True
        if(os.path.exists(fileName)):
            document = Document(fileName)
        else:
            document = Document()
            document.add_heading(u'领导信息', 0)

        document.add_heading(unicode(list[0].strip()+" "+list[1]+" "+list[2]),level=3)
        p = document.add_paragraph("")
        p.add_run(unicode(list[3])).italic = True
        p.add_run(unicode("    "+list[5]))
        document.add_paragraph(unicode("生年:  "+list[7]))
        document.add_paragraph(unicode("性别:  "+list[6]))
        document.add_paragraph(unicode("籍贯:  "+list[8]))
        document.add_paragraph(unicode("学历:  "+list[9]))
        document.add_paragraph(unicode("简历:  "+list[10]))
        document.add_paragraph(unicode(""))

        document.save(fileName)

'''
副省级领导抓取
'''
class viceProvince(object):
    def __init__(self):
        self.baseUrl = 'http://ldzl.people.com.cn/dfzlk/front/fusheng.htm'
        self.baseRoot = 'http://ldzl.people.com.cn/dfzlk/front/'
        self.flag = 0
        self.excelFileName = './data/viceProvince.xls'
        self.wordFileName = './data/officerWord/viceProvince.docx'


    def start(self):
        http = Http.Http()
        logging.info("++++++++++++副省级领导信息+++++++++++++")
        code, msg, res = http.open(self.baseUrl)
        if code == 200:
            print '++++++++++++reopen'
            soup = bs(res, 'html.parser')
            provinceList = soup.find_all('h2', class_='red2')
            for province in provinceList:
                provinceName = province.text
                for line in province.find_next_sibling('div', class_='ld2 new_ld').find_all('li', class_=False):
                    areaName = line.span.text
                    base = line.find_all('a')
                    secretaryBoss = base[1].string
                    secretaryBossUrl = base[1]['href']
                    govBoss = base[2].string
                    govBossUrl = base[2]['href']

                    partyBean = bean()
                    partyBean.setProvince(provinceName)
                    partyBean.setCity(areaName)
                    partyBean.setName(secretaryBoss)
                    if secretaryBoss!='空缺':
                        partyRes = county.resumeInfo(secretaryBossUrl,self.baseRoot)
                        if len(partyRes)!=0:
                            partyBean.setPosition(partyRes[0])
                            partyBean.setSex(partyRes[1])
                            partyBean.setBirth(partyRes[2])
                            partyBean.setNativePlace(partyRes[3])
                            partyBean.setEducation(partyRes[4])
                            partyBean.setResume(partyRes[5])

                    govBean = bean()
                    govBean.setProvince(provinceName)
                    govBean.setCity(areaName)
                    govBean.setName(govBoss)
                    if govBoss!='空缺':
                        govRes = county.resumeInfo(govBossUrl,self.baseRoot)
                        if len(govRes)!=0:
                            govBean.setPosition(govRes[0])
                            govBean.setSex(govRes[1])
                            govBean.setBirth(govRes[2])
                            govBean.setNativePlace(govRes[3])
                            govBean.setEducation(govRes[4])
                            govBean.setResume(govRes[5])
                    # print '省:'+partyBean.province+'  市、城区:'+partyBean.city+'  姓名:'+partyBean.name+'  职位:'+partyBean.position+'  性别:'+partyBean.sex+'  生日:'+partyBean.birth+'  籍贯:'+partyBean.nativePlace+'  学历:'+partyBean.education
                    # print '省:'+govBean.province+'  市、城区:'+govBean.city+'  姓名:'+govBean.name+'  职位:'+govBean.position+'  性别:'+govBean.sex+'  生日:'+govBean.birth+'  籍贯:'+govBean.nativePlace+'  学历:'+govBean.education

                    partyList = [partyBean.province,partyBean.city,partyBean.name,'党委书记',partyBean.position,partyBean.sex,partyBean.birth,partyBean.nativePlace,partyBean.education,partyBean.resume]
                    govList = [govBean.province,govBean.city,govBean.name,'政府',govBean.position,govBean.sex,govBean.birth,govBean.nativePlace,govBean.education,govBean.resume]
                    # self.flag = viceProvince.write2excel(self.flag, self.excelFileName, partyList)
                    # self.flag = viceProvince.write2excel(self.flag, self.excelFileName, govList)
                    # print '++++++++++写入至xls文件中的  '+str(self.flag-2)+'  和  '+str(self.flag-1)+'  行'
                    # common.write2Word(self.wordFileName,partyList)
                    # common.write2Word(self.wordFileName,govList)
                    # print "正在写入"
                    key = 2
                    officerDictParty = {
                        "province":partyBean.province,
                        "city":partyBean.city,
                        "name":partyBean.name,
                        "duty":partyList[3],
                        "position":partyBean.position,
                        "sex":partyBean.sex,
                        "birth":partyBean.birth,
                        "nativePlace":partyBean.nativePlace,
                        "education":partyBean.education,
                        "resume":partyBean.resume
                    }
                    officerDictGov = {
                        "province": govBean.province,
                        "city": govBean.city,
                        "name": govBean.name,
                        "duty": govList[3],
                        "position": govBean.position,
                        "sex": govBean.sex,
                        "birth": govBean.birth,
                        "nativePlace": govBean.nativePlace,
                        "education": govBean.education,
                        "resume": govBean.resume
                    }
                    if self.filterMongo(officerDictParty,key):
                        logging.warn("[+]数据更新\t"+partyBean.city+"\t"+partyBean.duty+"\t"+partyBean.name)
                        MongoDao.save(key,officerDictParty)
                    else:
                        logging.warn("[+]数据未更新\t" + partyBean.city + "\t" + partyBean.duty + "\t" + partyBean.name)

                    if self.filterMongo(officerDictGov,key):
                        logging.warn("[+]数据更新\t"+govBean.city+"\t"+govBean.duty+"\t"+govBean.name)
                        MongoDao.save(key,officerDictGov)
                    else:
                        logging.warn("[+]数据未更新\t"+govBean.city+"\t"+govBean.duty+"\t"+govBean.name)

        else:
            logging.error("[!]network error")
            exit(1)

    def filterMongo(self,officerDict,key):
        return MongoDao.getData(officerDict,key)

    @staticmethod
    def write2excel(flag, excelFileName, list=None):
        if flag == 0:
            if (os.path.exists(excelFileName)):
                os.remove(excelFileName)
            excelWorkBook = xlwt.Workbook()
            sheet = excelWorkBook.add_sheet('viceProvince')
            sheet.write(0, 0, u'省、直辖市')
            sheet.write(0, 1, u'市,直辖市城区')
            sheet.write(0, 2, u'姓名')
            sheet.write(0, 3, u'党委,政府')
            sheet.write(0, 4, u'职位')
            sheet.write(0, 5, u'性别')
            sheet.write(0, 6, u'生年')
            sheet.write(0, 7, u'籍贯')
            sheet.write(0, 8, u'学历')
            sheet.write(0, 9, u'简历')
            flag += 1
            excelWorkBook.save(excelFileName)
        else:
            import xlrd
            from xlutils.copy import copy
            rb = xlrd.open_workbook(excelFileName)
            wb = copy(rb)
            ws = wb.get_sheet(0)
            i = 0
            for data in list:
                ws.write(flag, i, unicode(data))
                i += 1
            flag += 1
            wb.save(excelFileName)
        return flag

'''
省级领导抓取
'''
class province(object):
    def __init__(self):
        self.baseUrl = 'http://ldzl.people.com.cn/dfzlk/front/personIndex.htm'
        self.baseRoot = 'http://ldzl.people.com.cn/dfzlk/front/'
        self.flag = 0
        self.excelFileName = './data/province.xls'
        self.wordFileName = './data/officerWord/province.docx'


    def start(self):
        logging.info("++++++++++++省级领导信息+++++++++++++")
        http = Http.Http()
        code, msg, res = http.open(self.baseUrl)
        if code == 200:
            logging.info("++++++++++++reopen")
            soup = bs(res, 'html.parser')
            provinceList = soup.find_all('h2', class_='red2')
            for province in provinceList:
                provinceName = province.text
                for line in province.find_next_sibling('ul', class_='clearfix ld new_ld2').find_all('dd', class_=False):
                    area = provinceName
                    duty = line.span.text
                    name = line.a.string
                    url = line.a['href']

                    officer = bean()
                    officer.setProvince(area)
                    officer.setName(name)
                    officer.setDuty(duty)

                    if name != '空缺':
                        govRes = county.resumeInfo(url, self.baseRoot)
                        if len(govRes) != 0:
                            officer.setPosition(govRes[0])
                            officer.setSex(govRes[1])
                            officer.setBirth(govRes[2])
                            officer.setNativePlace(govRes[3])
                            officer.setEducation(govRes[4])
                            officer.setResume(govRes[5])

                    officerList = [officer.province,officer.name,officer.duty,officer.position,officer.sex,officer.birth,officer.nativePlace,officer.education,officer.resume]
                    # self.flag = self.writeToexcel(self.flag,self.excelFileName,officerList)
                    # print name+'  已经写入 '+self.excelFileName+' 的 '+str(self.flag-1)+' 行'
                    # common.write2Word(self.wordFileName,officerList)
                    key = 1
                    officerDict = {
                        "province":officer.province,
                        "name":officer.name,
                        "duty":officer.duty,
                        "position":officer.position,
                        "sex":officer.sex,
                        "birth":officer.birth,
                        "nativePlace":officer.nativePlace,
                        "education":officer.education,
                        "resume":officer.resume
                    }
                    if self.filterMongo(officerDict,key):
                        logging.warn("[=]数据更新\t"+area+'\t'+duty+'\t'+name+'\t'+url)
                        MongoDao.save(key,officerDict)
                    else:
                        logging.warn("[=]数据未更新\t"+area+'\t'+duty+'\t'+name+'\t'+url)
        else:
            logging.warn("[-]network error")
            exit(1)

    def filterMongo(self,officerDict,key):
        return MongoDao.getData(officerDict,key)



    def writeToexcel(self,flag, excelFileName, list=None):
        if flag == 0:
            if (os.path.exists(excelFileName)):
                os.remove(excelFileName)
            excelWorkBook = xlwt.Workbook()
            sheet = excelWorkBook.add_sheet('province')
            sheet.write(0, 0, u'省、直辖市')
            sheet.write(0, 1, u'姓名')
            sheet.write(0, 2, u'职务')
            sheet.write(0, 3, u'职位')
            sheet.write(0, 4, u'性别')
            sheet.write(0, 5, u'生年')
            sheet.write(0, 6, u'籍贯')
            sheet.write(0, 7, u'学历')
            sheet.write(0, 8, u'简历')
            flag += 1
            excelWorkBook.save(excelFileName)
        else:
            import xlrd
            from xlutils.copy import copy
            rb = xlrd.open_workbook(excelFileName)
            wb = copy(rb)
            ws = wb.get_sheet(0)
            i = 0
            for data in list:
                ws.write(flag, i, unicode(data))
                i += 1
            flag += 1
            wb.save(excelFileName)
        return flag


class MongoDao(object):
    @staticmethod
    def save(key,officerDict):
        officerDict["key"] = key
        return mongoDb.insert(officerDict)

    @staticmethod
    def getData(officerDict,key):
        condition = {}
        if key==1:
            condition = {
                "province":officerDict["province"],
                "duty":officerDict["duty"]
            }
        if key==2:
            condition = {
                "city": officerDict["city"],
                "position": officerDict["position"]
            }
        if key==3:
            condition = {
                "city": officerDict['city'],
                "area": officerDict['area'],
                "position" : officerDict['position']
            }

        res = mongoDb.find_one(condition)
        if res is None or res["name"] != officerDict["name"]:
            return True

        if res["name"] == officerDict["name"]:
            return False




def test():
    resList = []
    soup = bs(open('./data/test.html'),"html.parser")
    postion =  soup.find('span',class_='red2').string
    if postion==None:
        postion=''

    resList.append(postion)
    infoBase = soup.find('dd').find('p')
    infoBaseParser = str(infoBase).split('\n')[:-1]
    for line in infoBaseParser:
        value = line.split('</b>')[1].split('<br>')[0]
        if value==None:
            value=''
        resList.append(value)

    info = soup.find('div',class_='p2j_text')
    resList.append(str(info.text))

    govRes = resList
    govCommitteePosition = govRes[0]
    govCommitteeSex = govRes[1]
    govCommitteeBirth = govRes[2]
    govCommitteeNativePlace = govRes[3]
    govCommitteeEducation = govRes[4]
    govCommitteeResume = govRes[5]

    print govCommitteePosition
    print govCommitteeSex
    print govCommitteeBirth
    print govCommitteeNativePlace
    print govCommitteeEducation

    print'  职位:' + govCommitteePosition + '  性别:' + govCommitteeSex + '  生年:' + govCommitteeBirth + '  籍贯:' + govCommitteeNativePlace + '  教育:' + govCommitteeEducation

def testWord():
    list = ["福建省","福州市","鼓楼区","杭东","","区委书记"]
    county.write2Word("./data/word.docx",list)
    county.write2Word("./data/word.docx",list)



if __name__ == '__main__':
    from util import Logging
    Logging.Logging("../log/renmin.log")

    # testWord()
    # vp = viceProvince()
    # vp.start()

    county = county()
    county.start()

    # province = province()
    # province.start()
