#!coding:utf-8
'''

支持腐败官员定期爬取

'''
import os
import sys
sys.path.append("..")
from bs4 import BeautifulSoup as bs
from docx import Document

from ztool import Http
import logging
from ztool import ZQueue
from ztool import const
from ztool import Db
import time
from util import Mongo as mg


saveToWordPath = '../data/societyTask/officerCorruptWord/'

mysqlDict = {
    "host":"localhost",
    "user":"root",
    "password":"shacha",
    "dbname":"officer"
}

mongoDb = mg.Mongo()


class corruptSpider(object):
    def __init__(self,url,level,type):
        self.url = url
        self.level = level
        self.type = type
        self.http = Http.Http()

    def start(self):
        http = self.http
        code,msg,res = http.open(self.url)
        if code==200:
            self.dataParser(res)
            logging.info("[+]over")
        else:
            raise Exception("network error")

    def dataParser(self,res):
        soup = bs(res,'html.parser')
        page = soup.find('div','page').text.strip().split('(')[1].split(',')[0]
        errorCount = 0
        urlList = []
        logging.info("[+]parser the html")
        queueStartTime = time.time()
        logging.info("[+]push queue")
        for i in range(int(page)):
            tempUrl = self.url
            if i!=0:
                tempUrl = self.url+'index_'+str(i)+".html"
            code,msg,res = self.http.open(tempUrl)
            if code==200:
                urlList.append(self.getChildUrl(res))
            else:
                logging.warn("[+]"+tempUrl+"  network error")
                errorCount += 1
                continue
        queueEndTime = time.time()
        logging.info('[+]queue time: ' + str(queueEndTime - queueStartTime)+'(s)\terror: '+str(errorCount))

        oneList = self.mergeList(urlList)
        filterList = self.filterRequest(oneList)
        self.getInfo(filterList)


    def filterRequest(self,oneList):
        latest = MongoDao.getLasted(self.level,self.type)
        resList = []
        if latest is not None:
            latestTime = str(latest[0]['time'])
            logging.info("[++]filter--->latestime:"+latestTime)
            titleList = []
            for i in latest:
                titleList.append(i['title'])

            for one in oneList:
                if one[2]>latestTime or (one[2]==latestTime and one[1] not in titleList):
                    resList.append(one)
            return resList
        else:
            return oneList

    def mergeList(self,urlList):
        oneList = []
        for a in urlList:
            for urlBean in a:
                oneList.append(urlBean)
        return oneList


    def getInfo(self,oneList):
        sumRequst = len(oneList)
        for k,request in enumerate(oneList):
            code,msg,res = self.http.open(self.url+request[0][1:])
            if code==200:
                logging.info("[++]"+'process: '+str(k)+'/'+str(sumRequst))
                resource,info = self.infoParser(res)
                b = bean(self.level,self.type,request[2]+' 00:00:00',request[1],info,self.url+request[0][1:],resource)
                MongoDao.save(b)
                # dao.saveToWord(b)
            else:
                logging.warn("[-]network error\t"+" error url:"+request[0][1:])
                continue


    def infoParser(self,res):
        soup = bs(res,'html.parser')
        resource = soup.find('em','e e1').text.split(u'：')[1].strip()
        infoBS = soup.find_all('div','TRS_Editor')
        if len(infoBS) == 1:
            info = infoBS[0].text
        else:
            info = infoBS[1].text
        return [resource,info]

    def getChildUrl(self,html):
        soup = bs(html,'html.parser')
        listElement = soup.find('ul','list_news_dl fixed')
        urlList = []
        for oneElement in listElement.find_all('li'):
            tempList = []
            # 链接
            tempList.append(oneElement.a['href'])
            # 标题
            tempList.append(oneElement.a.text.strip())
            # 时间点
            tempList.append(oneElement.span.text.strip())
            urlList.append(tempList)

        return urlList


class Mysqldao(object):
    @staticmethod
    def save(bean):
        sql = "insert into corrupt(level,type,time,title,content,resourceUrl,resource) VALUES (%d,%d,'%s','%s','%s','%s','%s')"%(bean.level,bean.type,bean.time,bean.title,bean.content,bean.resourceUrl,bean.resource)
        db = Db.Db(**mysqlDict)
        if db.insert(sql)==1:
            print '++++++success:'+bean.title

    @staticmethod
    def getLasted(level,type):
        sql = 'select * from corrupt where type=%d and level=%d group by time desc limit 1'%(type,level)
        db = Db.Db(**mysqlDict)
        res = db.execute(sql)
        if len(res)==0:
            return None
        return res

    @staticmethod
    def saveToWord(bean):
        levelList = ['zggb','sggb']
        typeList = ['zjsc','djcf']
        fileName = saveToWordPath + levelList[bean.level] + '_' + typeList[bean.type] + '.docx'
        try:
            if os.path.exists(fileName):
                document = Document(fileName)
            else:
                document = Document()
                document.add_heading(unicode(levelList[bean.level] + '_' + typeList[bean.type]), 0)

            document.add_heading(unicode(bean.title), level=3)
            p = document.add_paragraph("")
            p.add_run(unicode(bean.time+'      '+bean.resource+'         '+bean.resourceUrl)).italic = True
            resumeP = document.add_paragraph(unicode("内容："))
            resumeP.add_run(unicode(bean.content))
            document.save(fileName)
        except Exception:
            print 'fail to save to word'


class MongoDao(object):
    @staticmethod
    def save(bean):
        beanDict = bean.object2Dict()
        return mongoDb.db["corrupt"].insert_one(beanDict)


    @staticmethod
    def getLasted(level,type):
        res = mongoDb.db["corrupt"].find({"level":level,"type":type}).sort([
            ("time",-1)
        ]).limit(1)
        if res.count()==0:
            return None
        else:
            resList = []
            for i in res:
                resList.append(i)
            return resList





class bean(object):
    def __init__(self,level,type,time,title,content,resourceUrl,resource):
        self.level = level
        self.type = type
        self.time = time
        self.title = title
        self.content = content
        self.resourceUrl = resourceUrl
        self.resource = resource

    def object2Dict(self):
        return {
            "level":self.level,
            "type":self.type,
            "time":self.time,
            "title":self.title,
            "content":self.content,
            "resourceUrl":self.resourceUrl,
            "resource":self.resource
        }

if __name__=='__main__':
    from util import Logging
    Logging.Logging("../log/corrupt.log")

    c = corruptSpider('http://www.ccdi.gov.cn/jlsc/zggb/jlsc_zggb/',0,0)
    c.start()
    c = corruptSpider('http://www.ccdi.gov.cn/jlsc/zggb/djcf_zggb/', 0, 1)
    c.start()
    c = corruptSpider('http://www.ccdi.gov.cn/jlsc/sggb/jlsc_sggb/', 1, 0)
    c.start()
    c = corruptSpider('http://www.ccdi.gov.cn/jlsc/sggb/djcf_sggb/', 1, 1)
    c.start()