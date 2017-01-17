#!coding:utf-8

'''
搜索县级官员信息，来投入到mongodb中
百度百科搜索官员信息
'''
import os
import xlrd
from bs4 import BeautifulSoup as bs
import xlwt
from docx import Document
from multiprocessing import Process
import sys
from util import BaikeSpider as Baike
from util import Mongo as mg
import logging

class importExcel(object):
    def __init__(self):
        self.dataPath = '../data/officerExcel/'

    def load(self, fileName):
        list = []
        dataSet = set()
        if os.path.exists(self.dataPath) and os.path.isfile(self.dataPath + fileName):
            excelName = self.dataPath + fileName
            tag_row = 1
            book = xlrd.open_workbook(excelName, encoding_override='utf-8')
            sh = book.sheet_by_index(0)

            for rx in range(tag_row, sh.nrows):
                province = str((sh.row(rx))[1].value)
                city = str(sh.row(rx)[3].value)
                area = str(sh.row(rx)[5].value)
                name = str(sh.row(rx)[7].value)
                code = str(long(sh.row(rx)[4].value))
                # print province + ' ' + city + ' ' + area + ' ' + name
                if len(name) != 0 and name != 'N/A' and name != '' and name != ' ':
                    dataSet.add(code.strip()+":"+province.strip() + ':' + city.strip() + ':' + area.strip() + ':' + name.strip())
            return dataSet
        else:
            # log.warn("file path error")
            return dataSet


class exportBase(object):
    def __init__(self):
        pass

    def init(self,exportPath,exportName,data=None):
        pass

    def export(self):
        pass


class exportExcel(exportBase):
    def __init__(self):
        self.data = None
        self.exportPath = None
        self.exportName = None

    def init(self,exportPath,exportName,data=None):
        self.exportPath = exportPath
        self.exportName = exportName
        if data!=None:
            self.data = data

    def export(self,data):
        self.data = data
        fileName = self.exportPath + self.exportName
        if os.path.isdir(self.exportPath) is False:
            os.mkdir(self.exportPath)

        if os.path.exists(self.exportPath+self.exportName):
            import xlrd
            from xlutils.copy import copy
            filename = self.exportPath+self.exportName
            rb = xlrd.open_workbook(filename,encoding_override='utf-8')
            sh = rb.sheet_by_index(0)
            wb = copy(rb)
            ws = wb.get_sheet(0)
            flag = int(((sh.row(0))[0].value))
            i = 0
            for line in self.data[:4]:
                ws.write(flag, i, unicode(line))
                i += 1
            ws.write(0,0,flag+1)
            wb.save(filename)
        else:
            excelWorkBook = xlwt.Workbook()
            sheet = excelWorkBook.add_sheet('sheet1')
            count = 0
            for line in self.data:
                sheet.write(1,count,unicode(line))
                count += 1
            sheet.write(0,0,2)
            # print self.exportPath+self.exportName
            # exit(1)
            excelWorkBook.save(unicode(self.exportPath+self.exportName))

class exportWord(exportBase):
    def __init__(self):
        self.data = None
        self.exportPath = None
        self.exportName = None

    def init(self,exportPath,exportName,data=None):
        self.exportPath = exportPath
        self.exportName = exportName
        if data!=None:
            self.data = data

    def load(self,data):
        self.data = data

    def export(self,data):
        self.data = data
        fileName = self.exportPath+self.exportName
        if os.path.isdir(self.exportPath) is False:
            os.mkdir(self.exportPath)
        try:
            if os.path.exists(fileName):
                document = Document(fileName)
            else:
                document = Document()
                document.add_heading(unicode("县长信息"), 0)

            document.add_heading(unicode(self.data[0]+"  "+self.data[1]), level=2)
            p = document.add_paragraph("")
            p.add_run(unicode(self.data[2])).italic = True
            document.add_paragraph(unicode(self.data[3]))

            for k,v in self.data[5].items():
                document.add_paragraph(unicode(str(k)+": "+str(v)))
            resumeP = document.add_paragraph(unicode("简历："))
            for list in self.data[4]:
                resumeP.add_run(unicode(str(list)+"\n"))

            document.save(fileName)
        except Exception as e:
            print e


class mergeExcel(exportBase):
    def __init__(self):
        self.data = None
        self.exportPath = None
        self.exportName = None

    def init(self,exportPath,exportName,data=None):
        self.exportPath = exportPath
        self.exportName = exportName
        if data!=None:
            self.data = data

    def load(self,data):
        self.data = data

    def export(self,data):
        self.data = data
        fileName = self.exportPath+self.exportName
        if os.path.isdir(self.exportPath) is False:
            os.mkdir(self.exportPath)
        try:
            name = self.data[2]
            pairs = self.data[5]
            resume = self.data[4]
            # print name
            import xlrd
            from xlutils.copy import copy
            rb = xlrd.open_workbook(fileName, encoding_override='utf-8')
            sh = rb.sheet_by_index(0)
            wb = copy(rb)
            ws = wb.get_sheet(0)
            # flag = int(((sh.row(0))[0].value))
            flag = None
            for i in range(2,sh.nrows):
                strName = unicode(sh.row(i)[7].value)
                if str(unicode(strName)) == name:
                    flag = i
                    break
            if flag is not None:
                num = 21
                for k,v in pairs.items():
                    ws.write(flag,num,unicode(v))
                    num += 1
                for v in resume:
                    ws.write(flag,num,unicode(v))
                    num += 1
            wb.save(fileName)

        except Exception as e:
            print e


class exportMongo(exportBase):
    def __init__(self):
        self.mongoDb = mg.Mongo().db
        self.mongoDbCounty = self.mongoDb["county"]
        self.mongoDbPerson = self.mongoDb["person"]

        self.data = None
        self.flag = 0

    def init(self,exportPath,exportName,data=None):
        pass


    def export(self,data):
        self.data = data
        code = data[0]
        province = data[1]
        city = data[2]
        area = data[3]

        name = data[4]
        simpleInfo = data[5]
        complexInfoList = data[6]
        complexParamDict = data[7]
        personInfoList = data[4:]

        objectId = self.__saveToMongo(personInfoList)

        self.__findByMongo(objectId,code,name)

    def __findByMongo(self,objectId,code,name):
        res = self.mongoDbCounty.find_one(
            {
                # "province":province,
                # "city":city,
                "code":code
            }
        )
        if res is None:
            logging.error("[!]error\t"+"code:"+code)
            exit(1)

        keyName = "gov" if self.flag==0 else "party"

        if res.has_key(keyName):
            for key, value in enumerate(res[keyName]):
                if value["name"] == name:
                    res[keyName][key]["relationId"] = objectId
            self.mongoDbCounty.update({"code":code},res)
            logging.info("[+]related to the county collection\t"+"code:"+code)
        else:
            logging.warn("[+++]the key missing\t" + "code:" + code)
            exit(1)

    def __saveToMongo(self,personInfoList):
        paramDict = personInfoList[3]

        paramList = []
        for k,v in paramDict.items():
            paramList.append(k+":"+v)

        resumeList = []
        for k in personInfoList[2]:
            resumeList.append(k)

        dataDict = {
            "name" : personInfoList[0],
            "simpleInfo" : personInfoList[1],
            "resume" : resumeList,
            "paramDict" : paramList
        }
        # print dataDict["name"]
        # print dataDict["simpleInfo"]
        # print dataDict["resume"]
        # print dataDict["paramDict"]
        return self.mongoDbPerson.insert_one(dataDict).inserted_id


    #
    def setFlag(self,flag):
        self.flag = flag


if __name__ == '__main__':
    baikeSp = Baike.Baike()
    baikeSp.setLogging("../log/bs.log")

    # ew = mergeExcel()
    # ew.init("../data/officerBaikeWord/mergeToExcel/","party.xls")
    ew = exportMongo()
    ew.init("","")
    ew.setFlag(1)
    baikeSp.exportChoice(ew)

    ie = importExcel()
    data = ie.load("县委书记名单.xls")
    baikeSp.importData(data)

    baikeSp.start()

    ew.setFlag(0)
    baikeSp.exportChoice(ew)
    data = ie.load("县长的名单.xlsx")
    baikeSp.importData(data)
    baikeSp.start()